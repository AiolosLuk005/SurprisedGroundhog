from __future__ import annotations

import argparse
import csv
import os
import subprocess
import sys
from pathlib import Path

from docx import Document as DocxDocument

from core.normalize_base import NormalizeResult, register
from core.settings import SETTINGS

TIMEOUT = SETTINGS.get("normalize", {}).get("timeouts", {}).get("docx", 1800)


def process_docx(in_path: str, out_dir: Path) -> None:
    doc = DocxDocument(in_path)
    out_dir.mkdir(parents=True, exist_ok=True)
    md_lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            md_lines.append(text + "\n\n")
    md_path = out_dir / "document.md"
    md_path.write_text("".join(md_lines), encoding="utf-8")
    for idx, table in enumerate(doc.tables, start=1):
        rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
        csv_path = out_dir / f"table_{idx}.csv"
        with open(csv_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)
            writer.writerows(rows)


def cli() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()
    process_docx(args.input, Path(args.out))
    return 0


class DocxNormalizer:
    name = "docx"
    priority = 100
    exts = ["docx"]

    def can_handle(self, path: str) -> bool:
        return path.lower().endswith(".docx")

    def normalize(self, path: str, out_root: str) -> NormalizeResult:
        cmd = [sys.executable, str(Path(__file__).resolve()), "--input", path, "--out", out_root]
        env = os.environ.copy()
        env.setdefault("PYTHONPATH", str(Path(__file__).resolve().parents[2]))
        try:
            proc = subprocess.run(cmd, capture_output=True, text=True, timeout=TIMEOUT, env=env)
        except Exception as e:
            return NormalizeResult(False, message=str(e))
        if proc.returncode != 0:
            msg = (proc.stderr or proc.stdout or "")[-800:]
            return NormalizeResult(False, message=msg)
        md_paths = [str(p) for p in Path(out_root).glob("*.md")]
        csv_paths = [str(p) for p in Path(out_root).glob("*.csv")]
        sidecar = str(Path(out_root) / "sidecar.json") if (Path(out_root) / "sidecar.json").exists() else None
        return NormalizeResult(True, md_paths=md_paths, csv_paths=csv_paths, sidecar=sidecar)


register(DocxNormalizer())

if __name__ == "__main__":
    raise SystemExit(cli())
